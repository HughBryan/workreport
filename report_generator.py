from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Pt, RGBColor

def load_json(json_path):
    with open(json_path, "r") as f:
        return json.load(f)

def is_number(value):
    try:
        float(value)
        return True
    except Exception:
        return False

def format_currency(value, decimals=2):
    try:
        value = round(float(value), decimals)
        if decimals == 0:
            return "${:,.0f}".format(value)
        else:
            return "${:,.2f}".format(value)
    except Exception:
        return str(value)

def calculate_broker_fee(base, broker_fee_pct, commission_pct, commission_without_gst, fixed_broker_fee=0):

    # If we are adding a fixed fee.
    if fixed_broker_fee > 0:
        commission_adjustment = 0

        # If we are still doing commission - need to add the commission for insurers that don't do commissions in their quotes.
        if commission_without_gst == 0:
            commission_adjustment = (base*commission_pct/100)-commission_without_gst

            # don't bother if its only a few cents.
            if commission_adjustment < 1:
                commission_adjustment = 0

 
        # if we are adding fixed fee: brokerfee + commission shortfall
        return round(fixed_broker_fee+commission_adjustment, 2)
    try:
        commission_without_gst_val = float(commission_without_gst)
    except (ValueError, TypeError):
        commission_without_gst_val = 0
    commission_shortfall_pct = max(commission_pct - commission_without_gst_val / base * 100 if base else 0, 0)
    effective_broker_fee_pct = broker_fee_pct + commission_shortfall_pct
    return round(base * (effective_broker_fee_pct / 100.0), 2)

def enrich_insurer_quotes(quotes_dict, broker_fee_pct, commission_pct, associate_split=0, fixed_broker_fee=0):
    enriched = {}
    for insurer, quote in quotes_dict.items():
        base = float(quote.get("base", 0) or 0)
        total = float(quote.get("total", 0) or 0)
        commission_without_gst = quote.get("commission_without_gst", 0) or 0
        try:
            commission_without_gst_val = float(commission_without_gst)
        except Exception:
            commission_without_gst_val = 0
        broker_fee = calculate_broker_fee(base, broker_fee_pct, commission_pct, commission_without_gst_val, fixed_broker_fee)
        broker_gst = round(broker_fee * 0.1, 2)
        remuneration = round(commission_without_gst_val + broker_fee, 2)
        final_total = round(total + broker_fee + broker_gst, 2)
        sm_remuneration = round(remuneration * (associate_split / 100), 2)
        broker_remuneration = round(remuneration - sm_remuneration, 2)
        enriched_quote = dict(quote)
        enriched_quote["broker_fee"] = broker_fee
        enriched_quote["broker_gst"] = broker_gst
        enriched_quote["remuneration"] = remuneration
        enriched_quote["sm_remuneration"] = sm_remuneration
        enriched_quote["broker_remuneration"] = broker_remuneration
        enriched_quote["final_total"] = final_total
        enriched_quote["_final_total_numeric"] = final_total
        enriched_quote["insurer"] = insurer
        enriched[insurer] = enriched_quote
    return enriched

def find_recommended(enriched_quotes):
    min_total = float("inf")
    best = None
    for insurer, quote in enriched_quotes.items():
        total_val = quote.get("_final_total_numeric", float("inf"))
        if total_val and total_val < min_total:
            min_total = total_val
            best = quote
    if best:
        result = dict(best)
        result.pop("_final_total_numeric", None)
        return result
    return {}

def flatten_data_for_replace(data, broker_fee_pct, commission_pct, strata_manager, fixed_broker_fee=0):
    flat = {}
    for k, v in data.get("general_info", {}).items():
        flat[k] = v
    associate_split = data.get("associate_split", 0)
    enriched_quotes = enrich_insurer_quotes(data.get("Quotes", {}), broker_fee_pct, commission_pct, associate_split, fixed_broker_fee)
    for insurer, insurer_data in enriched_quotes.items():
        for field, value in insurer_data.items():
            if not field.startswith("_"):
                if is_number(value) and field not in ["uwgst", "uw", "uwgst_fee"]:
                    flat[f"{insurer}.{field}"] = format_currency(value)
                else:
                    flat[f"{insurer}.{field}"] = value
    recommended = find_recommended(enriched_quotes)
    for field, value in recommended.items():
        if is_number(value):
            flat[f"recommended.{field}"] = format_currency(value)
        else:
            flat[f"recommended.{field}"] = value
    flat["broker_fee_pct"] = f"{broker_fee_pct}%"
    flat["commission_pct"] = f"{commission_pct}%"
    flat["strata_manager"] = strata_manager
    return flat

def set_cell_background(cell, rgb_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_color)
    tcPr.append(shd)

def set_cell_bottom_border(cell, color="357ABD", size="12"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:color'), color)
    borders.append(bottom)

def ensure_landscape_section(doc):
    # Add a new section starting on a new page
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width
    return new_section

def insert_market_summary_table(doc, quotes, recommended_insurer, broker_fee_pct, commission_pct, associate_split, fixed_broker_fee=0):
    placeholder = "{{market_summary_table}}"
    # This is the full master list of known insurers (for missing entries)
    master_insurers = [
        "Axis", "CHU", "Flex", "Hutch", "IIS", "Longitude", "QUS", "SCI", "SUU"
    ]
    underwriters = {
        "Axis": "XL Insurance Company",
        "CHU": "QBE Insurance (Australia) Limited",
        "Flex": "QBE Insurance (Australia) Limited",
        "Hutch": "Certain Underwriters at Lloyds of London",
        "IIS": "Certain Underwriters at Lloyds of London",
        "Longitude": "Chubb Insurance Australia Limited",
        "QUS": "Certain Underwriters at Lloyds of London",
        "SCI": "Allianz Australia Insurance Limited",
        "SUU": "CGU Insurance Limited"
    }

    # Gather all unique insurer keys: those present in the quotes + those in the master list
    all_insurers = list({insurer for insurer in quotes} | set(master_insurers))
    all_insurers.sort()

    # Prepare data for table rows
    row_data = []
    for insurer in all_insurers:
        data = quotes.get(insurer)
        if data:
            enriched = enrich_insurer_quotes({insurer: data}, broker_fee_pct, commission_pct, associate_split, fixed_broker_fee)
            premium = enriched[insurer].get("final_total")
            comment = "Recommended" if insurer == recommended_insurer else ""
        else:
            premium = None
            comment = "Insurer did not respond in time"
        insurer_label = f"{insurer} – Underwritten by {underwriters.get(insurer, 'Unknown')}"
        row_data.append((insurer_label, format_currency(premium) if premium else "", comment))

    col_widths = [Inches(2.5), Inches(2.0), Inches(2.5)]

    for i, p in enumerate(doc.paragraphs):
        if placeholder in p.text:
            parent = p._element.getparent()
            idx = parent.index(p._element)
            parent.remove(p._element)
            tbl = doc.add_table(rows=1 + len(row_data), cols=3)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = True

            # Set up headers
            headers = ["Insurer / Underwriter", "Premium Payable", "Comment"]
            for col in range(3):
                cell = tbl.cell(0, col)
                cell.text = headers[col]
                set_cell_background(cell, "FFFFFF")
                set_cell_bottom_border(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                        r.font.name = "Futura Bk BT (Body)"
                        r.font.color.rgb = RGBColor(0, 0, 0)

            # Add dynamic rows
            for row_idx, (insurer_label, premium, comment) in enumerate(row_data, 1):
                color = "e9edf7" if row_idx % 2 == 1 else "FFFFFF"
                cells = [insurer_label, premium, comment]
                for col_idx, value in enumerate(cells):
                    cell = tbl.cell(row_idx, col_idx)
                    cell.text = value if value else ""
                    set_cell_background(cell, color)
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for r in p.runs:
                            r.font.size = Pt(9)
                            r.font.name = "Futura Bk BT (Body)"
            parent.insert(idx, tbl._element)
            break

def insert_comparison_table(doc, quotes, broker_fee_pct, commission_pct, associate_split, fixed_broker_fee=0):
    ensure_landscape_section(doc)
    placeholder = "{{comparison_table}}"
    insurer_list = list(quotes.keys())
    if not insurer_list:
        return
    first_features = next(iter(quotes.values())).get("features", {})
    feature_keys = list(first_features.keys())
    enriched_quotes = enrich_insurer_quotes(quotes, broker_fee_pct, commission_pct, associate_split, fixed_broker_fee)
    feature_keys.insert(0, "Total Premium")
    total_cols = 1 + len(insurer_list)
    column_width = Inches(9.0 / total_cols)
    for i, p in enumerate(doc.paragraphs):
        if placeholder in p.text:
            parent = p._element.getparent()
            idx = parent.index(p._element)
            parent.remove(p._element)
            tbl = doc.add_table(rows=1 + len(feature_keys), cols=total_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = True
            for col in range(total_cols):
                cell = tbl.cell(0, col)
                cell.width = column_width
                cell.text = "Common Policy Features" if col == 0 else insurer_list[col - 1]
                set_cell_background(cell, "FFFFFF")
                set_cell_bottom_border(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(11)
                        r.font.name = "Futura Bk BT"
                        r.font.color.rgb = RGBColor(0, 0, 0)
            for row_idx, key in enumerate(feature_keys, 1):
                color = "e9edf7" if row_idx % 2 == 1 else "FFFFFF"
                cell = tbl.cell(row_idx, 0)
                cell.width = column_width
                cell.text = key
                set_cell_background(cell, color)
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                        r.font.name = "Futura Bk BT"
                for col_idx, insurer in enumerate(insurer_list):
                    if key == "Total Premium":
                        val = enriched_quotes.get(insurer, {}).get("final_total", "-")
                        val = format_currency(val, 2) if is_number(val) else "-"
                    else:
                        val = quotes.get(insurer, {}).get("features", {}).get(key, "-")
                        if is_number(val):
                            val = format_currency(val, 0)
                    cell = tbl.cell(row_idx, col_idx + 1)
                    cell.width = column_width
                    cell.text = str(val) if val else "-"
                    set_cell_background(cell, color)
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for r in p.runs:
                            r.font.size = Pt(9)
                            r.font.name = "Futura Bk BT"
            parent.insert(idx, tbl._element)
            break

def insert_conditions_table(doc, quotes):
    ensure_landscape_section(doc)
    placeholder = "{{conditions_table}}"
    total_cols = 2
    col_widths = [Inches(2.5), Inches(6.5)]
    for i, p in enumerate(doc.paragraphs):
        if placeholder in p.text:
            parent = p._element.getparent()
            idx = parent.index(p._element)
            parent.remove(p._element)
            tbl = doc.add_table(rows=1 + len(quotes), cols=total_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = True
            headers = ["Insurer", "Conditions / Endorsements"]
            for col in range(total_cols):
                cell = tbl.cell(0, col)
                cell.width = col_widths[col]
                cell.text = headers[col]
                set_cell_background(cell, "FFFFFF")
                set_cell_bottom_border(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(11)
                        r.font.name = "Futura Bk BT"
                        r.font.color.rgb = RGBColor(0, 0, 0)
            for row_idx, (insurer, quote) in enumerate(quotes.items(), 1):
                color = "e9edf7" if row_idx % 2 == 1 else "FFFFFF"
                values = [insurer, quote.get("conditions_or_endorsements", "-")]
                for col_idx, value in enumerate(values):
                    cell = tbl.cell(row_idx, col_idx)
                    cell.width = col_widths[col_idx]
                    cell.text = value if value else "-"
                    set_cell_background(cell, color)
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for r in p.runs:
                            r.font.size = Pt(10)
                            r.font.name = "Futura Bk BT"
            parent.insert(idx, tbl._element)
            break


def insert_invoice_comparison_table(doc, recommended_quote, previous_invoice_data=None):
    placeholder = "{{invoice_comparison_table}}"
    ensure_landscape_section(doc)

    row_labels = [
        "Insurer / Underwriter",
        "Base Premium",
        "ESL or FSL",
        "GST",
        "Stamp Duty",
        "Insurer / Underwriter Fee",
        "Insurer / Underwriter Fee GST",
        "Broker Fee",
        "Broker Fee GST",
        "Total Insurance Premium"
    ]

    quote_keys = [
        "insurer",
        "base",
        "esl",
        "gst",
        "stamp",
        "underwriter_fee",
        "underwriter_fee_gst",
        "broker_fee",
        "broker_gst",
        "final_total"
    ]

    invoice_keys = [
        "insurer_/_underwriter",
        "base_premium",
        "esl",
        "gst",
        "stamp_duty",
        "insurer_/_underwriter_fee",
        "insurer_/_underwriter_gst",
        "broker_fee",
        "broker_fee_gst",
        "total_premium"
    ]

    col_headers = ["Itemised Insurance Costs"]
    if previous_invoice_data:
        col_headers.append("Last Policy Period")
    col_headers.append("This Policy Period")

    num_cols = len(col_headers)
    col_width = Inches(9.0 / num_cols)

    for para in doc.paragraphs:
        if placeholder in para.text:
            parent = para._element.getparent()
            idx = parent.index(para._element)
            parent.remove(para._element)

            tbl = doc.add_table(rows=1 + len(row_labels), cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.autofit = True

            # Set headers
            for col_idx, text in enumerate(col_headers):
                cell = tbl.cell(0, col_idx)
                cell.width = col_width
                cell.text = text
                set_cell_background(cell, "FFFFFF")
                set_cell_bottom_border(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(11)
                        r.font.name = "Futura Bk BT"
                        r.font.color.rgb = RGBColor(0, 0, 0)

            # Fill rows
            for row_idx, label in enumerate(row_labels, 1):
                color = "e9edf7" if row_idx % 2 == 1 else "FFFFFF"
                row = tbl.row_cells(row_idx)

                is_last_row = row_idx == len(row_labels)

                # First column: static label
                row[0].text = label
                set_cell_background(row[0], color)
                for p in row[0].paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(9)
                        r.font.name = "Futura Bk BT"

                col_offset = 1

                # Second column: previous invoice
                if previous_invoice_data:
                    key = invoice_keys[row_idx - 1]
                    val = previous_invoice_data.get(key, "-")
                    if isinstance(val, (int, float)):
                        val = format_currency(val)
                    row[col_offset].text = val or "-"
                    set_cell_background(row[col_offset], color)
                    for p in row[col_offset].paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for r in p.runs:
                            r.font.size = Pt(9)
                            r.font.name = "Futura Bk BT"
                            if is_last_row:
                                r.font.bold = True
                    col_offset += 1

                # Third column: recommended quote
                rec_key = quote_keys[row_idx - 1]
                val = recommended_quote.get(rec_key, "-")
                if is_number(val):
                    val = format_currency(val)
                row[col_offset].text = val or "-"
                set_cell_background(row[col_offset], color)
                for p in row[col_offset].paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = "Futura Bk BT"
                        if is_last_row:
                            r.font.bold = True

            if row_idx == len(row_labels):
                for cell in tbl.row_cells(row_idx - 1):
                    set_cell_bottom_border(cell)

            parent.insert(idx, tbl._element)
            break

from docx.shared import Pt
from docx.shared import Pt, RGBColor

def insert_disclosure_section(doc, strata_manager):
    placeholder = "{{Disclosure}}"

    # If strata_manager is missing or blank, just clear the placeholder
    if not strata_manager:
        for para in doc.paragraphs:
            if placeholder in para.text:
                para.text = ""
        return

    disclosure_heading = "DISCLOSURE OF COMMERCIAL AGREEMENT"
    disclosure_paragraphs = [
        "[The broking services outlined in this report are provided by:",
        "Clearlake Insurance Brokers ACN 651 113 861.",
        f"Clearlake Insurance Brokers and {strata_manager} have a referral relationship/distribution relationship. Both parties may receive income as a result of placing your annual insurance and providing ongoing support. The full details of the remuneration payable are detailed in the *Itemised Insurance Costs* section of this document.]"
    ]

    for para in doc.paragraphs:
        if placeholder in para.text:
            parent = para._element.getparent()
            idx = parent.index(para._element)
            parent.remove(para._element)

            # Title paragraph (blue)
            title_para = doc.add_paragraph()
            run = title_para.add_run(disclosure_heading)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Futura Bk BT"
            run.font.color.rgb = RGBColor(0, 112, 192)

            # Body paragraphs
            for text in disclosure_paragraphs:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.name = "Futura Bk BT (Body)"

            # Insert into correct location
            parent.insert(idx, title_para._element)
            for i in range(len(disclosure_paragraphs)):
                parent.insert(idx + i + 1, doc.paragraphs[-(len(disclosure_paragraphs) - i)]._element)
            break


def generate_report(template_path, output_path, data, broker_fee_pct, commission_pct, associate_split, strata_manager, fixed_broker_fee=0, previous_invoice_data=None):
    doc = Document(template_path)
    data["associate_split"] = associate_split
    data["strata_manager"] = strata_manager
    replace_dict = flatten_data_for_replace(data, broker_fee_pct, commission_pct, strata_manager, fixed_broker_fee)

    bold_keys = {
        "recommended.insurer",
        "recommended.final_total",
        "recommended.remuneration"
    }

    # Replace in paragraphs
    for para in doc.paragraphs:
        if "{{expiring_premium}}" in para.text:
            para.clear()
            if previous_invoice_data and "total_premium" in previous_invoice_data:
                val = previous_invoice_data["total_premium"]
                formatted = format_currency(val)
                run1 = para.add_run("The expiring premium (exclusive of any adjustments throughout the year) was ")
                run2 = para.add_run(formatted)
                run2.bold = True


        for key, value in replace_dict.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                # Save existing text
                text_parts = para.text.split(placeholder)
                para.clear()
                run = para.add_run(text_parts[0])
                run.bold = False

                # Insert bold value if it's in the bold_keys set
                bold_run = para.add_run(str(value))
                bold_run.bold = key in bold_keys

                run2 = para.add_run(text_parts[1])
                run2.bold = False

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replace_dict.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        for para in cell.paragraphs:
                            if placeholder in para.text:
                                parts = para.text.split(placeholder)
                                para.clear()
                                para.add_run(parts[0])
                                bold_run = para.add_run(str(value))
                                bold_run.bold = key in bold_keys
                                para.add_run(parts[1] if len(parts) > 1 else "")

    enriched_quotes = enrich_insurer_quotes(data.get("Quotes", {}), broker_fee_pct, commission_pct, associate_split, fixed_broker_fee)
    recommended = find_recommended(enriched_quotes)
    insert_comparison_table(doc, data.get("Quotes", {}), broker_fee_pct, commission_pct, associate_split, fixed_broker_fee)
    insert_conditions_table(doc, data.get("Quotes", {}))
    insert_market_summary_table(doc, data.get("Quotes", {}), recommended.get("insurer"), broker_fee_pct, commission_pct, associate_split, fixed_broker_fee)
    insert_invoice_comparison_table(doc, recommended, previous_invoice_data)
    insert_disclosure_section(doc,strata_manager)



    # Save output
    doc.save(output_path)
